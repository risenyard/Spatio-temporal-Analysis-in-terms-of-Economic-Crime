import docx
import os
import re
import pandas as pd

#正则1：名字 性别 出生年（必有）
regex1 = r"被告人(.*)，([男女]).*(19..)年"
regex12 = r"被告人(.*)，([男女]).*(200.)年"
#正则1-2：民族
regex2 = r"被告人.*(.{1,4}族)"
#正则1-3：文化
regex3 = r"被告人.*(..)文化"
regex32= r"(文盲)"
#正则2：判决年
regex4 = r"(20..).*川.*刑.*号"
#正则2-1：判罪 刑期
#regex3 = r"被告人.*犯(.*罪)，判处(.*)，并处罚金人民币(.*)元"
regex5 = r"被告人.*犯(.*罪)，判处(.{2,20})[；|，]并处罚金人民币(.{1,10})元"
regex52 = r"被告人.*犯(.*罪)，判处(.{2,20})[；|，]并处罚金(.{1,10})元"
regex6 = r"被告人.*(无罪)"

#dataframe数组
df=pd.DataFrame(columns=('名字','性别','出生年','判决年','民族','文化','判罪','刑期','罚金'))


#获取文档和结果
lis = os.walk(r"Verdict-of-all-cities\巴中")          #获取文档对象
for path, dir_list, file_list in lis:
    for file_name in file_list:
      try:
        a = '';b = '';c = '';d = '';e = '';f = '';g = '';h='';z='' #初始化参数
        pa = os.path.join(path, file_name) ;print(pa)
        file = docx.Document(pa)  #读入文档
        con = 0   #标记初始化
        for i in range(len(file.paragraphs)):  # 基本正则（正则1）
            # print("第"+str(i)+"段："+file.paragraphs[i].text)
            matchObj1 = re.search(regex1, file.paragraphs[i].text)
            matchObj12 = re.search(regex12, file.paragraphs[i].text)
            if matchObj1:
                con = con + 1
                a = matchObj1.group(1);  # print(matchObj1.group(1))
                b = matchObj1.group(2);  # print(matchObj1.group(2))
                c = matchObj1.group(3);  # print(matchObj1.group(3))

                matchObj2 = re.search(regex2, file.paragraphs[i].text)
                if matchObj2:  # 正则1-2民族
                    d = matchObj2.group(1);  # print(matchObj2.group(1))
                else:
                    d = '';# print('no people')

                matchObj3 = re.search(regex3, file.paragraphs[i].text)
                if matchObj3:  # 正则1-3文化
                    e = matchObj3.group(1);  # print(matchObj3.group(1))
                else:
                    e = '';# print('no study')
                    matchObj32 = re.search(regex32, file.paragraphs[i].text)
                    if matchObj32:  # 文盲可能
                        e = matchObj32.group(1);  # print(matchObj3.group(1))

                for k in range(len(file.paragraphs)):
                    matchObj4 = re.search(regex4, file.paragraphs[k].text)
                    if matchObj4:  # 正则2：判决年
                        z = matchObj4.group(1)
                        break
                    else:
                        z = ''

                con2 = 1  # 标记
                for j in range(len(file.paragraphs)):  # 正则3：罪名-刑期-罚金
                    matchObj5 = re.search(regex5, file.paragraphs[j].text)
                    if matchObj5:
                        if con == con2:
                            f = matchObj5.group(1);  # print(matchObj5.group(1))
                            g = matchObj5.group(2);  # print(matchObj5.group(2))
                            h = matchObj5.group(3);  # print(matchObj5.group(3))
                            break
                        con2 = con2 + 1
                    else:
                        matchObj52 = re.search(regex52, file.paragraphs[j].text)
                        if matchObj52:  # 模式52可能
                            if con == con2:
                                f = matchObj52.group(1);  # print(matchObj5.group(1))
                                g = matchObj52.group(2);  # print(matchObj5.group(2))
                                h = matchObj52.group(3);  # print(matchObj5.group(3))
                                break
                            con2 = con2 + 1

                df = df.append({'名字': a, '性别': b, '出生年': c,'判决年': z, '民族': d, '文化': e, '判罪': f, '刑期': g, '罚金': h},
                               ignore_index=True)


            elif matchObj12:
                matchObj1=matchObj12
                con = con + 1
                a = matchObj1.group(1);  # print(matchObj1.group(1))
                b = matchObj1.group(2);  # print(matchObj1.group(2))
                c = matchObj1.group(3);  # print(matchObj1.group(3))

                matchObj2 = re.search(regex2, file.paragraphs[i].text)
                if matchObj2:  # 民族
                    d = matchObj2.group(1);  # print(matchObj2.group(1))
                else:
                    d = '';# print('no people')

                matchObj3 = re.search(regex3, file.paragraphs[i].text)
                if matchObj3:  # 文化
                    e = matchObj3.group(1);  # print(matchObj3.group(1))
                else:
                    e = '';# print('no study')
                    matchObj32 = re.search(regex32, file.paragraphs[i].text)
                    if matchObj32:  # 文盲可能
                        e = matchObj32.group(1);  # print(matchObj3.group(1))

                for k in range(len(file.paragraphs)):
                    matchObj4 = re.search(regex4, file.paragraphs[k].text)
                    if matchObj4:  # 判决年
                        z = matchObj4.group(1)
                        break
                    else:
                        z = ''

                con2 = 1  # 标记
                for j in range(len(file.paragraphs)):
                    matchObj5 = re.search(regex5, file.paragraphs[j].text)
                    if matchObj5:
                        if con == con2:
                            f = matchObj5.group(1);  # print(matchObj5.group(1))
                            g = matchObj5.group(2);  # print(matchObj5.group(2))
                            h = matchObj5.group(3);  # print(matchObj5.group(3))
                            break
                        con2 = con2 + 1
                    else:
                        matchObj52 = re.search(regex52, file.paragraphs[j].text)
                        if matchObj52:  # 模式52可能
                            if con == con2:
                                f = matchObj52.group(1);  # print(matchObj5.group(1))
                                g = matchObj52.group(2);  # print(matchObj5.group(2))
                                h = matchObj52.group(3);  # print(matchObj5.group(3))
                                break
                            con2 = con2 + 1

                df = df.append({'名字': a, '性别': b, '出生年': c,'判决年': z, '民族': d, '文化': e, '判罪': f, '刑期': g, '罚金': h},
                               ignore_index=True)



      except:
        print("oops")
        continue

print(df)

#df.to_excel('test.xlsx')